# document_parser.py - Multi-format document parser
import re
import fitz  # PyMuPDF
from pathlib import Path
from html.parser import HTMLParser
from dataclasses import dataclass

import yaml
from docx import Document as DocxDocument


@dataclass
class ParsedDocument:
    """Output of the parsing stage."""
    text: str
    metadata: dict
    pages: list  # Per-page text for PDFs, single-element list for others


# ---------------------------------------------------------------------------
# Metadata extraction helpers
# ---------------------------------------------------------------------------

def _extract_yaml_front_matter(text: str) -> tuple[dict, str]:
    """Extract YAML front matter from markdown text. Returns (metadata, body)."""
    match = re.match(r"^---\s*\n(.*?)\n---\s*\n", text, re.DOTALL)
    if not match:
        return {}, text
    try:
        meta = yaml.safe_load(match.group(1)) or {}
    except yaml.YAMLError:
        return {}, text
    body = text[match.end():]
    return meta, body


class _HTMLMetaExtractor(HTMLParser):
    """Extract <meta> tag attributes from HTML."""
    def __init__(self):
        super().__init__()
        self.meta = {}

    def handle_starttag(self, tag, attrs):
        if tag == "meta":
            attrs_dict = dict(attrs)
            name = attrs_dict.get("name", "")
            content = attrs_dict.get("content", "")
            if name and content:
                self.meta[name] = content


class _HTMLTextExtractor(HTMLParser):
    def __init__(self):
        super().__init__()
        self._pieces = []
        self._skip = False

    def handle_starttag(self, tag, attrs):
        if tag in ("script", "style"):
            self._skip = True

    def handle_endtag(self, tag):
        if tag in ("script", "style"):
            self._skip = False
        if tag in ("p", "div", "br", "h1", "h2", "h3", "h4", "li", "tr", "section"):
            self._pieces.append("\n")

    def handle_data(self, data):
        if not self._skip:
            self._pieces.append(data)

    def get_text(self):
        return "".join(self._pieces).strip()


# ---------------------------------------------------------------------------
# Parsers
# ---------------------------------------------------------------------------

def parse_pdf_pymupdf(file_path: Path) -> ParsedDocument:
    """Parse a PDF using PyMuPDF with OCR fallback for scanned pages."""
    doc = fitz.open(str(file_path))
    pages = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text()

        # If the page produced very little text, try OCR
        if len(text.strip()) < 50:
            text = page.get_text("text", flags=fitz.TEXT_PRESERVE_WHITESPACE)
            if len(text.strip()) < 50:
                try:
                    tp = page.get_textpage_ocr(
                        language="eng",
                        dpi=300,
                        full=True,
                    )
                    text = page.get_text("text", textpage=tp)
                except Exception:
                    text = ""  # OCR failed; log and continue

        pages.append(text)

    doc.close()
    full_text = "\n\n".join(pages)

    return ParsedDocument(
        text=full_text,
        metadata={
            "source": file_path.name,
            "file_path": str(file_path),
            "doc_type": "pdf",
            "total_pages": len(pages),
        },
        pages=pages,
    )


def parse_html(file_path: Path) -> ParsedDocument:
    """Parse an HTML file by stripping tags and extracting <meta> metadata."""
    raw = file_path.read_text(encoding="utf-8", errors="replace")

    # Extract metadata from <meta> tags
    meta_extractor = _HTMLMetaExtractor()
    meta_extractor.feed(raw)
    front_matter = meta_extractor.meta

    # Extract visible text
    extractor = _HTMLTextExtractor()
    extractor.feed(raw)
    text = extractor.get_text()

    metadata = {
        "source": file_path.name,
        "file_path": str(file_path),
        "doc_type": "html",
        "total_pages": 1,
    }
    # Merge front-matter fields (title, category, doc_type, etc.)
    for key in ("title", "category", "doc_type", "last_updated", "owner", "classification"):
        if key in front_matter:
            metadata[key] = front_matter[key]

    return ParsedDocument(text=text, metadata=metadata, pages=[text])


def parse_markdown(file_path: Path) -> ParsedDocument:
    """Parse a markdown file, extracting YAML front matter as metadata."""
    raw = file_path.read_text(encoding="utf-8", errors="replace")
    front_matter, body = _extract_yaml_front_matter(raw)

    metadata = {
        "source": file_path.name,
        "file_path": str(file_path),
        "doc_type": "markdown",
        "total_pages": 1,
    }
    # Merge front-matter fields
    for key in ("title", "category", "doc_type", "last_updated", "owner", "classification"):
        if key in front_matter:
            metadata[key] = front_matter[key]

    return ParsedDocument(text=body, metadata=metadata, pages=[body])


def parse_docx(file_path: Path) -> ParsedDocument:
    """Parse a DOCX file using python-docx."""
    doc = DocxDocument(str(file_path))
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    return ParsedDocument(
        text=text,
        metadata={
            "source": file_path.name,
            "file_path": str(file_path),
            "doc_type": "docx",
            "total_pages": 1,
        },
        pages=[text],
    )


def parse_document(file_path: Path) -> ParsedDocument:
    """Route a file to the appropriate parser."""
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf_pymupdf(file_path)
    elif suffix in (".html", ".htm"):
        return parse_html(file_path)
    elif suffix == ".md":
        return parse_markdown(file_path)
    elif suffix == ".docx":
        return parse_docx(file_path)
    elif suffix == ".txt":
        return parse_markdown(file_path)  # treat as plain text
    else:
        raise ValueError(f"Unsupported file type: {suffix}")
